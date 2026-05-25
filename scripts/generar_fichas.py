"""
Genera un documento Word con fichas de revisión por tipo de acto canónico.
Una ficha por tipo, agrupadas por categoría, listas para UAT.
"""

import csv
from collections import defaultdict
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

INPUT_CSV  = r"C:\Users\Florencia Martinez\notarial\scripts\resultado_consolidado.csv"
OUTPUT_DOC = r"C:\Users\Florencia Martinez\notarial\scripts\fichas_modelos_notariales.docx"

CALIDAD_COLOR = {
    "alta":  RGBColor(0x2E, 0x86, 0x48),   # verde
    "media": RGBColor(0xB7, 0x7A, 0x00),   # amarillo oscuro
    "baja":  RGBColor(0xC0, 0x39, 0x2B),   # rojo
    "":      RGBColor(0x88, 0x88, 0x88),
}

CALIDAD_LABEL = {"alta": "LISTA", "media": "REVISAR", "baja": "REQUIERE TRABAJO", "": "SIN DATOS"}

ORDEN_CATEGORIAS = [
    "Transmisión de dominio",
    "Garantías reales",
    "Poderes y mandatos",
    "Contratos",
    "Actas notariales",
    "Certificaciones notariales",
    "Certificaciones registrales",
    "Sociedades",
    "Sucesiones",
    "Documentos administrativos",
    "Otros",
]

def set_cell_bg(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)

def agregar_fila_tabla(tabla, label, valor, bold_label=True):
    row = tabla.add_row()
    c0, c1 = row.cells
    p0 = c0.paragraphs[0]
    r0 = p0.add_run(label)
    r0.bold = bold_label
    r0.font.size = Pt(9)
    r0.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
    set_cell_bg(c0, "F5F5F5")
    p1 = c1.paragraphs[0]
    r1 = p1.add_run(str(valor) if valor else "—")
    r1.font.size = Pt(9.5)

def generar():
    rows = list(csv.DictReader(open(INPUT_CSV, encoding="utf-8-sig")))

    por_categoria = defaultdict(list)
    for r in rows:
        por_categoria[r["categoria"]].append(r)

    doc = Document()

    # Márgenes
    for section in doc.sections:
        section.top_margin    = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    # Portada
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("CATÁLOGO DE MODELOS NOTARIALES")
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(0x1A, 0x23, 0x32)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run("Revisión preliminar — Uso interno")
    r2.font.size = Pt(11)
    r2.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    doc.add_paragraph()

    # Índice
    p_idx = doc.add_paragraph()
    r_idx = p_idx.add_run("RESUMEN POR CATEGORÍA")
    r_idx.bold = True
    r_idx.font.size = Pt(11)

    tabla_idx = doc.add_table(rows=1, cols=3)
    tabla_idx.style = "Table Grid"
    hdr = tabla_idx.rows[0].cells
    for c, t in zip(hdr, ["Categoría", "Tipos", "Listos"]):
        hdr[0].paragraphs[0].runs[0] if hdr[0].paragraphs[0].runs else None
        p = c.paragraphs[0]
        r = p.add_run(t)
        r.bold = True
        r.font.size = Pt(9)
        set_cell_bg(c, "1A2332")
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    for cat in ORDEN_CATEGORIAS:
        grupo = por_categoria.get(cat, [])
        if not grupo:
            continue
        listos = sum(1 for r in grupo if r["usar_como_preset"] == "si")
        row = tabla_idx.add_row()
        row.cells[0].paragraphs[0].add_run(cat).font.size = Pt(9)
        row.cells[1].paragraphs[0].add_run(str(len(grupo))).font.size = Pt(9)
        row.cells[2].paragraphs[0].add_run(str(listos)).font.size = Pt(9)

    doc.add_page_break()

    ficha_num = 0
    for cat in ORDEN_CATEGORIAS:
        grupo = por_categoria.get(cat, [])
        if not grupo:
            continue

        # Encabezado de categoría
        p_cat = doc.add_paragraph()
        r_cat = p_cat.add_run(f"▌ {cat.upper()}")
        r_cat.bold = True
        r_cat.font.size = Pt(13)
        r_cat.font.color.rgb = RGBColor(0x3A, 0x7C, 0xA5)
        doc.add_paragraph()

        for row in sorted(grupo, key=lambda x: x["tipo_canonico"]):
            ficha_num += 1
            calidad = row.get("calidad_preset", "")
            color   = CALIDAD_COLOR.get(calidad, CALIDAD_COLOR[""])
            label   = CALIDAD_LABEL.get(calidad, "")

            # Título de la ficha
            p_titulo = doc.add_paragraph()
            r_num  = p_titulo.add_run(f"{ficha_num:02d}. ")
            r_num.bold = True
            r_num.font.size = Pt(12)
            r_num.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)
            r_tipo = p_titulo.add_run(row["tipo_canonico"])
            r_tipo.bold = True
            r_tipo.font.size = Pt(12)
            r_tipo.font.color.rgb = RGBColor(0x1A, 0x23, 0x32)
            r_est = p_titulo.add_run(f"   [{label}]")
            r_est.bold = True
            r_est.font.size = Pt(9)
            r_est.font.color.rgb = color

            # Tabla de datos
            tabla = doc.add_table(rows=0, cols=2)
            tabla.style = "Table Grid"
            tabla.columns[0].width = Cm(4)
            tabla.columns[1].width = Cm(12)

            variantes = int(row.get("variantes_en_grupo", 1))
            agregar_fila_tabla(tabla, "Variantes en archivo", f"{variantes} documento/s analizado/s")
            agregar_fila_tabla(tabla, "Archivo representativo", row.get("archivo", ""))
            agregar_fila_tabla(tabla, "Año", row.get("anio", ""))
            agregar_fila_tabla(tabla, "Vigente CCyC", "Sí" if row.get("vigente_ccyc") == "True" else "No")
            agregar_fila_tabla(tabla, "Variables detectadas", row.get("variables_detectadas", ""))
            agregar_fila_tabla(tabla, "Observaciones", row.get("observaciones", ""))

            # Espacio para notas manuales
            p_notas = doc.add_paragraph()
            r_notas = p_notas.add_run("Notas de revisión: _________________________________________________________________")
            r_notas.font.size = Pt(9)
            r_notas.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)

            doc.add_paragraph()

        doc.add_page_break()

    doc.save(OUTPUT_DOC)
    print(f"Word generado: {OUTPUT_DOC}")
    print(f"Total fichas : {ficha_num}")

if __name__ == "__main__":
    generar()
